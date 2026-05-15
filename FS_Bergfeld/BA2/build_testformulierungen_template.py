from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = "Testformulierungen_Beobachtungsassistent_Tiefenstruktur.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_paragraph(paragraph, size=9, bold=False, color="1B2A32"):
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(cell, text, size=9, bold=False, color="1B2A32"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def main():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    title = document.add_paragraph()
    title_run = title.add_run("Testformulierungen für den Beobachtungsassistenten Tiefenstruktur")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(27, 42, 50)
    title.paragraph_format.space_after = Pt(4)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(
        "Bitte knapp und realistisch formulieren: so, wie du es im Unterricht tatsächlich tippen würdest. "
        "Eine Zeile = ein möglicher Testfall. Richtung ist optional."
    )
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(82, 98, 107)
    subtitle.paragraph_format.space_after = Pt(10)

    rows = []
    contexts = [
        ("Einstieg", "Plenum", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Kurzmarker"]),
        ("Erarbeitung", "Unterrichtsgespräch", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Arbeitsphase", "Einzelarbeit", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Arbeitsphase", "Partnerarbeit", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Arbeitsphase", "Gruppenarbeit", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Gruppendiskussion", "Gruppenarbeit", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Sicherung", "Plenum", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Feedback", "Einzelarbeit / Gruppe", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Übergang", "Plenum / Gruppenarbeit", ["Lehrerzitat", "Schülerzitat", "Beobachtungsnotiz", "Kurzmarker", "Beobachtungsnotiz"]),
        ("Freie fachliche Beobachtung", "beliebig", ["freie Beobachtung", "freie Beobachtung", "freie Beobachtung", "Kurzmarker", "Notiz"]),
    ]
    for phase, social_form, types in contexts:
        for kind in types:
            rows.append((phase, social_form, kind))

    table = document.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths = [0.35, 1.25, 1.45, 1.18, 5.5, 1.05]
    headers = ["Nr.", "Phase", "Sozialform", "Typ", "Formulierung / Zitat / Kurzmarker", "Richtung optional"]
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for index, cell in enumerate(hdr.cells):
        set_cell_width(cell, widths[index])
        set_cell_margins(cell)
        set_cell_shading(cell, "DDEAF3")
        add_text(cell, headers[index], size=8.5, bold=True, color="1B2A32")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for i, (phase, social_form, kind) in enumerate(rows, start=1):
        row = table.add_row()
        values = [str(i), phase, social_form, kind, "", "grün / blau / frei / ?"]
        for col, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col])
            set_cell_margins(cell, top=100, bottom=100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i % 5 == 1:
                set_cell_shading(cell, "F6F9FB")
            else:
                set_cell_shading(cell, "FFFFFF")
            add_text(cell, values[col], size=8.5 if col != 4 else 9, bold=False, color="1B2A32" if col != 5 else "6A767D")

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note_run = note.add_run(
        "Hinweis: Die Spalte „Richtung“ ist freiwillig. Wenn du sie offen lässt, teste ich nur die Item-Zuordnung. "
        "Wenn du grün/blau/frei/? einträgst, prüfe ich zusätzlich, ob die Tendenz passt."
    )
    note_run.font.name = "Arial"
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(82, 98, 107)

    document.save(OUTPUT)


if __name__ == "__main__":
    main()
